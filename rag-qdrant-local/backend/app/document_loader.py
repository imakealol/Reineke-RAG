"""Extract text from supported document types into a uniform record format.

A `LoadedSegment` is the smallest unit returned by a loader: a piece of text
plus *positional* metadata (page, sheet, row range). Chunkers consume these
segments downstream to build the actual Qdrant chunks.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

from .office_converter import (
    OfficeConversionError,
    convert_doc_to_docx,
    convert_xls_to_xlsx,
)
from .utils import get_logger

log = get_logger(__name__)


class DocumentLoadError(RuntimeError):
    pass


class RequiresOCRError(DocumentLoadError):
    """Raised for image-only PDFs that have no extractable text."""


@dataclass
class LoadedSegment:
    """One coherent piece of text from a source document."""

    text: str
    document_type: str  # "pdf" | "docx" | "xlsx"
    page: Optional[int] = None
    sheet: Optional[str] = None
    row_start: Optional[int] = None
    row_end: Optional[int] = None
    has_header: bool = False
    extras: dict = field(default_factory=dict)


@dataclass
class LoadedDocument:
    file_path: Path
    file_name: str
    file_extension: str
    document_type: str
    segments: List[LoadedSegment]


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _load_pdf(path: Path) -> List[LoadedSegment]:
    from pypdf import PdfReader  # local import keeps cold-start light

    try:
        reader = PdfReader(str(path))
    except Exception as exc:  # pragma: no cover — defensive
        raise DocumentLoadError(f"Could not open PDF: {exc}") from exc

    segments: List[LoadedSegment] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            log.warning("Page %s of %s: extract_text failed (%s)", i, path.name, exc)
            text = ""
        text = text.strip()
        if not text:
            continue
        segments.append(
            LoadedSegment(text=text, document_type="pdf", page=i)
        )

    if not segments:
        raise RequiresOCRError(
            f"PDF '{path.name}' has no extractable text — likely a scan. "
            f"Mark as requires_ocr."
        )
    return segments


# ---------------------------------------------------------------------------
# DOCX (and DOC via LibreOffice)
# ---------------------------------------------------------------------------

def _table_to_markdown(table) -> str:
    rows: List[List[str]] = []
    for row in table.rows:
        rows.append([(cell.text or "").strip().replace("\n", " ") for cell in row.cells])

    if not rows:
        return ""

    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]

    header = rows[0]
    body = rows[1:] if len(rows) > 1 else []

    lines = ["| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join(["---"] * width) + " |")
    for r in body:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def _load_docx(path: Path) -> List[LoadedSegment]:
    from docx import Document  # python-docx

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise DocumentLoadError(f"Could not open DOCX: {exc}") from exc

    segments: List[LoadedSegment] = []

    paragraphs: List[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            paragraphs.append(t)

    if paragraphs:
        # Concatenate paragraphs separated by blank lines so the chunker
        # can split on natural boundaries.
        segments.append(
            LoadedSegment(
                text="\n\n".join(paragraphs),
                document_type="docx",
            )
        )

    for idx, table in enumerate(doc.tables, start=1):
        md = _table_to_markdown(table)
        if md:
            segments.append(
                LoadedSegment(
                    text=f"[Table {idx}]\n{md}",
                    document_type="docx",
                    extras={"table_index": idx},
                )
            )

    if not segments:
        raise DocumentLoadError(f"DOCX '{path.name}' produced no extractable text.")
    return segments


def _load_doc_legacy(path: Path) -> List[LoadedSegment]:
    converted = convert_doc_to_docx(path)
    return _load_docx(converted)


# ---------------------------------------------------------------------------
# XLSX (and XLS via LibreOffice)
# ---------------------------------------------------------------------------

def _looks_like_header(row: List[str]) -> bool:
    """Heuristic: a row is a header if every cell is a non-empty short string
    and at most one cell parses as a number."""
    if not row or not all(isinstance(c, str) and c.strip() for c in row):
        return False
    numeric = 0
    for c in row:
        try:
            float(c.replace(",", "."))
            numeric += 1
        except ValueError:
            pass
    return numeric <= 1 and all(len(c) <= 64 for c in row)


def _row_to_str(row) -> List[str]:
    return ["" if v is None else str(v) for v in row]


def _load_xlsx(path: Path) -> List[LoadedSegment]:
    from openpyxl import load_workbook

    try:
        wb = load_workbook(filename=str(path), data_only=True, read_only=True)
    except Exception as exc:
        raise DocumentLoadError(f"Could not open XLSX: {exc}") from exc

    segments: List[LoadedSegment] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_iter = ws.iter_rows(values_only=True)
        rows: List[List[str]] = []
        for r in rows_iter:
            row_str = _row_to_str(r)
            if any(c.strip() for c in row_str):
                rows.append(row_str)

        if not rows:
            continue

        header_row: Optional[List[str]] = None
        body_rows = rows
        if _looks_like_header(rows[0]):
            header_row = rows[0]
            body_rows = rows[1:]

        # Body row index 0 corresponds to spreadsheet row 2 if header present.
        first_data_row_number = 2 if header_row is not None else 1

        if not body_rows:
            # Header-only sheet — emit it as one tiny segment so users can
            # still query for the column names.
            if header_row:
                segments.append(
                    LoadedSegment(
                        text=_format_table([header_row]),
                        document_type="xlsx",
                        sheet=sheet_name,
                        row_start=1,
                        row_end=1,
                        has_header=True,
                    )
                )
            continue

        segments.append(
            LoadedSegment(
                text=_format_table(([header_row] if header_row else []) + body_rows),
                document_type="xlsx",
                sheet=sheet_name,
                row_start=first_data_row_number,
                row_end=first_data_row_number + len(body_rows) - 1,
                has_header=header_row is not None,
                extras={"header_row": header_row} if header_row else {},
            )
        )

    wb.close()

    if not segments:
        raise DocumentLoadError(f"XLSX '{path.name}' produced no extractable text.")
    return segments


def _load_xls_legacy(path: Path) -> List[LoadedSegment]:
    converted = convert_xls_to_xlsx(path)
    return _load_xlsx(converted)


def _format_table(rows: List[List[str]]) -> str:
    """Render rows as TSV — compact and embedding-friendly."""
    return "\n".join("\t".join(r) for r in rows)


# ---------------------------------------------------------------------------
# Public dispatch
# ---------------------------------------------------------------------------

_DISPATCH = {
    ".pdf": (_load_pdf, "pdf"),
    ".docx": (_load_docx, "docx"),
    ".doc": (_load_doc_legacy, "docx"),
    ".xlsx": (_load_xlsx, "xlsx"),
    ".xls": (_load_xls_legacy, "xlsx"),
}


def load_document(path: Path) -> LoadedDocument:
    ext = path.suffix.lower()
    if ext not in _DISPATCH:
        raise DocumentLoadError(f"Unsupported file type: {ext}")

    loader, doctype = _DISPATCH[ext]
    try:
        segments = loader(path)
    except OfficeConversionError as exc:
        raise DocumentLoadError(str(exc)) from exc

    return LoadedDocument(
        file_path=path,
        file_name=path.name,
        file_extension=ext,
        document_type=doctype,
        segments=segments,
    )
