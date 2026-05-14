"""Extractor smoke tests for DOCX, XLSX, and HTML (legacy .doc/.xls require LibreOffice)."""

from pathlib import Path

import pytest

from app.document_loader import DocumentLoadError, load_document


def _build_xlsx(path: Path) -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Server"
    ws.append(["Hostname", "IP", "Role"])
    ws.append(["srv-01", "10.0.0.1", "DB"])
    ws.append(["srv-02", "10.0.0.2", "App"])
    ws.append(["srv-03", "10.0.0.3", "Web"])
    wb.create_sheet("Empty")
    wb.save(path)


def _build_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Test", level=1)
    doc.add_paragraph("Dies ist ein Absatz mit Inhalt.")
    doc.add_paragraph("Zweiter Absatz mit weiteren Informationen.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Spalte A"
    table.rows[0].cells[1].text = "Spalte B"
    table.rows[1].cells[0].text = "Wert 1"
    table.rows[1].cells[1].text = "Wert 2"
    doc.save(path)


def test_xlsx_extraction(sandbox_root: Path):
    p = sandbox_root / "test.xlsx"
    _build_xlsx(p)

    loaded = load_document(p)
    assert loaded.document_type == "xlsx"

    server_seg = next(s for s in loaded.segments if s.sheet == "Server")
    assert server_seg.has_header is True
    assert server_seg.row_start == 2
    assert server_seg.row_end == 4
    assert "srv-01" in server_seg.text
    assert "Hostname" in server_seg.text


def test_docx_extraction(sandbox_root: Path):
    p = sandbox_root / "test.docx"
    _build_docx(p)

    loaded = load_document(p)
    assert loaded.document_type == "docx"
    full = "\n".join(s.text for s in loaded.segments)
    assert "Absatz" in full
    assert "Spalte A" in full
    assert "Wert 1" in full


def test_unsupported_extension_raises(sandbox_root: Path):
    p = sandbox_root / "foo.txt"
    p.write_text("hello")
    with pytest.raises(DocumentLoadError):
        load_document(p)


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------

_HTML_SAMPLE_PROSE = """<!DOCTYPE html>
<html lang="de">
<head>
    <meta charset="utf-8">
    <title>Versuchsprotokoll 1576 Lubrizol</title>
    <style>body { color: red; }</style>
    <script>alert('xss');</script>
</head>
<body>
    <nav>Navigation: <a href="/home">Home</a></nav>
    <header><h2>Sitewide Header</h2></header>
    <main>
        <h1>Versuch 1576 Lubrizol AM 400-2591</h1>
        <p>Verwendetes Aggregat: Konusmischer Typ AM 400-2591.</p>
        <ul>
            <li>Charge: 250 kg</li>
            <li>Dauer: 45 Minuten</li>
        </ul>
        <blockquote>Wichtige Anmerkung zur Probenahme.</blockquote>
    </main>
    <footer>(c) 2026 Reineke</footer>
</body>
</html>"""

_HTML_SAMPLE_WITH_TABLE = """<!DOCTYPE html>
<html><body>
<h1>Messwerte</h1>
<p>Übersicht der erfassten Werte.</p>
<table>
    <thead>
        <tr><th>Zeitpunkt</th><th>Temperatur</th><th>Druck</th></tr>
    </thead>
    <tbody>
        <tr><td>10:00</td><td>22 °C</td><td>1.0 bar</td></tr>
        <tr><td>10:15</td><td>45 °C</td><td>1.2 bar</td></tr>
    </tbody>
</table>
</body></html>"""

_HTML_SAMPLE_NOISE_ONLY = """<html><head><title>n</title>
<script>var x = 1;</script><style>p{}</style></head>
<body><nav>menu</nav><footer>foot</footer></body></html>"""


def test_html_extraction_prose_and_strips_noise(sandbox_root: Path):
    p = sandbox_root / "versuch.html"
    p.write_text(_HTML_SAMPLE_PROSE, encoding="utf-8")

    loaded = load_document(p)

    assert loaded.document_type == "html"
    assert loaded.file_extension == ".html"
    assert len(loaded.segments) == 1
    seg = loaded.segments[0]
    assert seg.document_type == "html"

    # Title + main content survives
    assert "Versuchsprotokoll 1576 Lubrizol" in seg.text  # <title>
    assert "Versuch 1576 Lubrizol AM 400-2591" in seg.text  # <h1>
    assert "Konusmischer Typ AM 400-2591" in seg.text  # <p>
    assert "Charge: 250 kg" in seg.text  # <li>
    assert "Wichtige Anmerkung" in seg.text  # <blockquote>

    # Noise stripped — scripts, styles, nav and footer must not leak through
    assert "alert(" not in seg.text
    assert "color: red" not in seg.text
    assert "Navigation:" not in seg.text
    assert "Sitewide Header" not in seg.text
    assert "(c) 2026 Reineke" not in seg.text

    # Title preserved in extras for downstream display
    assert seg.extras.get("title") == "Versuchsprotokoll 1576 Lubrizol"


def test_html_extraction_with_table(sandbox_root: Path):
    p = sandbox_root / "messwerte.htm"
    p.write_text(_HTML_SAMPLE_WITH_TABLE, encoding="utf-8")

    loaded = load_document(p)

    # Both .html and .htm dispatch to the same loader
    assert loaded.document_type == "html"
    assert loaded.file_extension == ".htm"

    # Prose segment + one table segment
    assert len(loaded.segments) == 2

    prose = next(s for s in loaded.segments if "table_index" not in s.extras)
    table_seg = next(s for s in loaded.segments if s.extras.get("table_index") == 1)

    # Prose carries the heading + intro paragraph
    assert "Messwerte" in prose.text
    assert "Übersicht der erfassten Werte" in prose.text
    # Table data must NOT appear in the prose segment (avoid duplication)
    assert "10:15" not in prose.text

    # Table renders as markdown with the data preserved
    assert "| Zeitpunkt | Temperatur | Druck |" in table_seg.text
    assert "| 10:00 | 22 °C | 1.0 bar |" in table_seg.text
    assert "| 10:15 | 45 °C | 1.2 bar |" in table_seg.text


def test_html_extraction_empty_content_raises(sandbox_root: Path):
    """A page that contains only chrome (nav/footer/scripts) must fail loudly
    so it is marked as `failed` in the documents table — not silently indexed
    as an empty document."""
    p = sandbox_root / "noise.html"
    p.write_text(_HTML_SAMPLE_NOISE_ONLY, encoding="utf-8")

    with pytest.raises(DocumentLoadError):
        load_document(p)


def test_html_extraction_handles_legacy_encoding(sandbox_root: Path):
    """German umlauts in a Windows-1252 / ISO-8859-1 file must round-trip."""
    p = sandbox_root / "legacy.html"
    body = (
        "<html><head><meta http-equiv='Content-Type' "
        "content='text/html; charset=iso-8859-1'><title>Legacy</title></head>"
        "<body><p>Größe und Maß für Schüttgüter.</p></body></html>"
    )
    p.write_bytes(body.encode("iso-8859-1"))

    loaded = load_document(p)
    text = "\n".join(s.text for s in loaded.segments)
    assert "Größe und Maß für Schüttgüter." in text
